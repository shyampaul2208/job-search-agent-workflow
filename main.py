from dotenv import load_dotenv
load_dotenv()
import os
import json
import uuid
from datetime import datetime

import streamlit as st
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.schema import Document
from langchain_groq import ChatGroq
from langgraph.graph import StateGraph, END

import docx
from PyPDF2 import PdfReader
from typing import TypedDict, List, Dict, Optional

# Optional CUDA fallback for embeddings
try:
    import torch  # type: ignore
    _EMBED_DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
except Exception:
    _EMBED_DEVICE = "cpu"

# ------------------------
# API keys
# ------------------------
groq_api_key = os.getenv("GROQ_API_KEY")
hf_api_key = os.getenv("HUGGINGFACEHUB_API_TOKEN")

# Storage paths
DATA_DIR = os.path.join(os.path.dirname(__file__), "data")
APPLICATIONS_PATH = os.path.join(DATA_DIR, "applications.json")
JOBS_PATH = os.path.join(DATA_DIR, "jobs.json")

os.makedirs(DATA_DIR, exist_ok=True)

# ------------------------
# Resume Parser
# ------------------------
def parse_resume(file):
    """Parse resume file into clean text."""
    text = ""
    if file.name.endswith(".pdf"):
        pdf = PdfReader(file)
        for page in pdf.pages:
            text += page.extract_text() or ""
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:  # fallback: txt or unknown
        text = file.read().decode("utf-8", errors="ignore")
    return text.strip()


def persist_application(application: Dict) -> None:
    """Append an application record to JSON storage."""
    try:
        if os.path.exists(APPLICATIONS_PATH):
            with open(APPLICATIONS_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
        else:
            data = []
    except Exception:
        data = []
    data.append(application)
    with open(APPLICATIONS_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_jobs() -> List[Dict]:
    """Load job postings from JSON file; if missing, return empty list."""
    try:
        if os.path.exists(JOBS_PATH):
            with open(JOBS_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return []

# ------------------------
# Sample Data
# ------------------------
job_postings = load_jobs()

# ------------------------
# Build FAISS Vector DB
# ------------------------
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2",
    cache_folder="./hf_cache",
    model_kwargs={"device": _EMBED_DEVICE}
)

def _job_to_content(j: Dict) -> str:
    skills = ", ".join(j.get("skills", []))
    remote = j.get("remote")
    remote_str = "remote" if remote is True else ("hybrid" if str(remote).lower() == "hybrid" else "onsite")
    return (
        f"Title: {j['title']}\n"
        f"Company: {j.get('company', '')}\n"
        f"Description: {j['description']}\n"
        f"Location: {j.get('location', '')}\n"
        f"Work Mode: {remote_str}\n"
        f"Skills: {skills}"
    )

docs = [
    Document(page_content=_job_to_content(j), metadata=j)
    for j in job_postings
]
if docs:
    vectorstore = FAISS.from_documents(docs, embeddings)
    retriever = vectorstore.as_retriever()
else:
    vectorstore = None
    retriever = None

############################
# LangGraph Workflow (Conversational)
############################

class AgentState(TypedDict):
    messages: List[Dict[str, str]]
    preferences: Dict[str, str]
    candidate_jobs: List[Dict]
    selected_job: Optional[Dict]
    resume_text: Optional[str]
    profile: Dict[str, str]
    answers: Dict[str, str]
    question_index: int
    stage: str


# Setup Groq LLM
llm = ChatGroq(model="moonshotai/kimi-k2-instruct-0905", api_key=groq_api_key)


def extract_preferences_with_llm(user_text: str) -> Dict[str, str]:
    """Use the LLM to extract coarse preferences from user text.
    Returns keys: role, skills, location, remote (as 'true'/'false'/'hybrid'/'' )."""
    try:
        sys = (
            "Extract job search preferences from the user's message. "
            "Return strict JSON with keys: role, skills (comma-separated), location, remote. "
            "Remote should be one of: true, false, hybrid, ''."
        )
        prompt = (
            f"SYSTEM:\n{sys}\n\nUSER:\n{user_text}\n\n"
            "Return JSON only."
        )
        resp = llm.invoke(prompt)
        content = resp.content.strip()
        # Find first JSON object in output
        start = content.find("{")
        end = content.rfind("}")
        if start != -1 and end != -1 and end > start:
            import json as _json
            parsed = _json.loads(content[start:end+1])
            # Normalize
            role = str(parsed.get("role", "")).strip()
            skills = str(parsed.get("skills", "")).strip()
            location = str(parsed.get("location", "")).strip()
            remote = str(parsed.get("remote", "")).strip().lower()
            return {
                "role": role,
                "skills": skills,
                "location": location,
                "remote": remote,
            }
    except Exception:
        pass
    return {"role": "", "skills": "", "location": "", "remote": ""}


def format_jobs_for_assistant(jobs: List[Dict]) -> str:
    lines = []
    for idx, j in enumerate(jobs, start=1):
        skills = ", ".join(j.get("skills", []))
        lines.append(
            f"{idx}. {j['title']} @ {j.get('company','')} — {j.get('location','')} — Skills: {skills}\n   {j['description']}"
        )
    return "\n\n".join(lines)


def match_job_from_user_text(user_text: str, candidates: List[Dict]) -> Optional[Dict]:
    text = user_text.lower()
    # Try by explicit index
    for idx, _ in enumerate(candidates, start=1):
        if str(idx) in text.split():
            return candidates[idx - 1]
    # Try by title substring
    for j in candidates:
        if j["title"].lower() in text:
            return j
    # Try by job id mention
    for j in candidates:
        if j["id"].lower() in text:
            return j
    return None


def router(state: AgentState) -> Dict:
    return state


def route_next(state: AgentState) -> str:
    messages = state.get("messages", [])
    last_role = messages[-1]["role"] if messages else None
    stage = state.get("stage", "greeting")

    if not messages:
        return "greet"

    if stage == "greeting":
        return "greet"

    if state.get("selected_job") is None:
        # Either retrieve or allow selection
        # If we already have candidates, try to match selection on user turn
        if state.get("candidate_jobs") and last_role == "user":
            return "select_job"
        # Otherwise only retrieve if the user asked about jobs
        if last_role == "user":
            user_text = messages[-1]["content"]
            if is_job_search_intent(user_text):
                return "retrieve"
            else:
                return "clarify"
        return "noop"

    # Job selected, wait for resume/profile
    if state.get("selected_job") and not state.get("resume_text"):
        # Streamlit UI will collect resume/profile; nothing to do until provided
        return "noop"

    # Resume provided -> ask or record Q&A
    if state.get("selected_job") and state.get("resume_text"):
        # If no question has been asked yet, start Q&A immediately
        if int(state.get("question_index", 0)) == 0:
            return "record_or_ask"
        # If last was assistant, we're waiting for user's answer
        if last_role == "assistant":
            return "noop"
        # If user responded, record and proceed
        return "record_or_ask"

    return "noop"


def greet(state: AgentState) -> Dict:
    greeting_text = (
        "👋 Hello! I can help you discover relevant jobs and apply.\n"
        "Tell me about the role, skills, and location you prefer (e.g., 'Remote Python role in India')."
    )
    return {
        "messages": state["messages"] + [{"role": "assistant", "content": greeting_text}],
        "stage": "job_search",
    }


def retrieve(state: AgentState) -> Dict:
    user_text = state["messages"][-1]["content"]
    prefs = extract_preferences_with_llm(user_text)
    merged_prefs = {**state.get("preferences", {}), **{k: v for k, v in prefs.items() if v}}

    # Build query text
    query = user_text
    if merged_prefs:
        query += "\n" + ", ".join([f"{k}:{v}" for k, v in merged_prefs.items() if v])

    if retriever is not None:
        results = retriever.get_relevant_documents(query)
        jobs = [r.metadata for r in results]
    else:
        jobs = job_postings
    if not jobs:
        jobs = job_postings  # fallback to all

    # Role-based filtering: if role specified or role keyword present (e.g., "software")
    desired_role = (merged_prefs.get("role") or "").lower()
    if not desired_role:
        lt = user_text.lower()
        # Simple keyword mapping
        if "software" in lt:
            desired_role = "software"
        elif "backend" in lt:
            desired_role = "backend"
        elif "frontend" in lt:
            desired_role = "frontend"
        elif "data" in lt:
            desired_role = "data"

    if desired_role:
        def _match_role(j: Dict) -> bool:
            title = j.get("title", "").lower()
            return desired_role in title
        filtered = [j for j in jobs if _match_role(j)]
        # If we asked for software, ensure only software jobs are shown
        if filtered:
            jobs = filtered

    body = (
        "Here are relevant jobs based on your preferences. "
        "Reply with the job number or title to proceed.\n\n"
        + format_jobs_for_assistant(jobs[:5])
    )

    return {
        "messages": state["messages"] + [{"role": "assistant", "content": body}],
        "preferences": merged_prefs,
        "candidate_jobs": jobs[:5],
        "stage": "job_search",
    }


def select_job(state: AgentState) -> Dict:
    user_text = state["messages"][-1]["content"]
    candidates = state.get("candidate_jobs", [])
    selected = match_job_from_user_text(user_text, candidates)
    if not selected:
        # Ask to clarify
        return {
            "messages": state["messages"] + [{"role": "assistant", "content": "Please choose a job by number or title from the list above."}],
            "stage": "job_search",
        }

    ask = (
        f"Great choice! ✅ You selected: {selected['title']}.\n\n"
        "Please upload your resume below and provide your name, email, and preferred location. "
        "Once submitted, I'll ask a few screening questions."
    )
    return {
        "messages": state["messages"] + [{"role": "assistant", "content": ask}],
        "selected_job": selected,
        "stage": "awaiting_resume",
    }


def record_or_ask(state: AgentState) -> Dict:
    job = state["selected_job"]
    q_index = int(state.get("question_index", 0))
    questions = job.get("questions", [])

    # If we have just received a user answer for the previous question
    if state["messages"] and state["messages"][-1]["role"] == "user" and q_index > 0:
        prev_q = questions[q_index - 1]
        user_ans = state["messages"][-1]["content"]
        new_answers = {**state.get("answers", {}), prev_q: user_ans}
    else:
        new_answers = state.get("answers", {})

    # If more questions remain, ask next
    if q_index < len(questions):
        next_q = questions[q_index]
        return {
            "messages": state["messages"] + [{"role": "assistant", "content": f"Question {q_index+1}/{len(questions)}: {next_q}"}],
            "answers": new_answers,
            "question_index": q_index + 1,
            "stage": "asking_questions",
        }

    # Otherwise finalize application
    application = {
        "id": str(uuid.uuid4()),
        "created_at": datetime.utcnow().isoformat() + "Z",
        "job_id": job.get("id"),
        "job_title": job.get("title"),
        "preferences": state.get("preferences", {}),
        "profile": state.get("profile", {}),
        "answers": new_answers,
        "resume_text": state.get("resume_text", ""),
        # "conversation": state.get("messages", []),
    }
    persist_application(application)

    done_msg = (
        "🎉 Your application has been submitted successfully! We'll be in touch.\n"
        f"Role: {job.get('title')}"
    )
    return {
        "messages": state["messages"] + [{"role": "assistant", "content": done_msg}],
        "answers": new_answers,
        "stage": "completed",
    }


def noop(state: AgentState) -> Dict:
    return state


def is_job_search_intent(text: str) -> bool:
    t = (text or "").lower()
    keywords = [
        "job", "jobs", "role", "position", "opening", "vacancy",
        "hiring", "apply", "list jobs", "find jobs", "opportunity",
        "engineer", "developer", "data scientist", "remote",
    ]
    if any(k in t for k in keywords):
        return True
    # Optional: LLM fallback classification
    if groq_api_key:
        try:
            prompt = (
                "Decide if the user is asking to discover jobs. "
                "Answer strictly 'yes' or 'no'.\n\nUser: " + text
            )
            resp = llm.invoke(prompt)
            ans = (resp.content or "").strip().lower()
            return ans.startswith("y")
        except Exception:
            return False
    return False


def clarify(state: AgentState) -> Dict:
    msg = (
        "I can help you find jobs. "
        "Tell me your preferred role, key skills, location, and whether you want remote, "
        "or simply say something like: 'Show remote Python jobs in India'."
    )
    return {
        "messages": state["messages"] + [{"role": "assistant", "content": msg}],
        "stage": "job_search",
    }


workflow = StateGraph(AgentState)
workflow.add_node("router", router)
workflow.add_node("greet", greet)
workflow.add_node("retrieve", retrieve)
workflow.add_node("select_job", select_job)
workflow.add_node("record_or_ask", record_or_ask)
workflow.add_node("noop", noop)
workflow.add_node("clarify", clarify)

workflow.set_entry_point("router")
workflow.add_conditional_edges(
    "router",
    route_next,
    {
        "greet": "greet",
        "retrieve": "retrieve",
        "select_job": "select_job",
        "record_or_ask": "record_or_ask",
        "clarify": "clarify",
        "noop": "noop",
    },
)
workflow.add_edge("greet", "router")
workflow.add_edge("retrieve", "router")
workflow.add_edge("select_job", "router")
workflow.add_edge("record_or_ask", "router")
workflow.add_edge("clarify", "router")
workflow.add_edge("noop", END)

app = workflow.compile()

############################
# Streamlit App (Conversational UI)
############################

st.set_page_config(page_title="Job Assistant", page_icon="💼", layout="centered")

custom_css = """
<style>
.stChatFloatingInputContainer { bottom: 0 !important; }
.assistant-msg { background: #F5F7FB; padding: 12px 14px; border-radius: 10px; border: 1px solid #E6EAF2; }
.user-msg { background: #E8FFF3; padding: 12px 14px; border-radius: 10px; border: 1px solid #CFEFDD; }
.meta-box { font-size: 12px; color: #5f6c7b; }
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

st.title("💼 Job Agent")
if groq_api_key is None:
    st.warning("GROQ_API_KEY not found. Set it in your environment for LLM responses.")

if "state" not in st.session_state:
    st.session_state.state = {
        "messages": [],
        "preferences": {},
        "candidate_jobs": [],
        "selected_job": None,
        "resume_text": None,
        "profile": {},
        "answers": {},
        "question_index": 0,
        "stage": "greeting",
    }

# On first load, produce a greeting
if not st.session_state.state["messages"]:
    st.session_state.state = app.invoke(st.session_state.state)

# Render conversation so far
for msg in st.session_state.state["messages"]:
    role = msg.get("role", "assistant")
    with st.chat_message(role):
        st.markdown(f"<div class='{'assistant-msg' if role=='assistant' else 'user-msg'}'>{msg['content']}</div>", unsafe_allow_html=True)

# UI: If awaiting resume/profile after selection
if st.session_state.state.get("selected_job") and not st.session_state.state.get("resume_text"):
    job = st.session_state.state["selected_job"]
    st.subheader(f"📄 Apply to {job['title']}")
    with st.form("resume_profile_form", clear_on_submit=False):
        name = st.text_input("Full Name", value=st.session_state.state.get("profile", {}).get("name", ""))
        email = st.text_input("Email", value=st.session_state.state.get("profile", {}).get("email", ""))
        pref_loc = st.text_input("Preferred Location", value=st.session_state.state.get("profile", {}).get("preferred_location", ""))
        resume_file = st.file_uploader("Upload Resume (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"], key="resume_file")
        submitted = st.form_submit_button("Submit Resume & Profile")
        if submitted:
            resume_text = None
            if resume_file is not None:
                resume_text = parse_resume(resume_file)
            if resume_text:
                st.session_state.state["resume_text"] = resume_text
                st.session_state.state["profile"] = {"name": name, "email": email, "preferred_location": pref_loc}
                # Trigger first question
                st.session_state.state = app.invoke(st.session_state.state)
                st.rerun()
            else:
                st.error("Please upload a valid resume file.")

# Chat input
user_input = st.chat_input("Describe jobs you want, or answer the question…")
if user_input:
    st.session_state.state["messages"].append({"role": "user", "content": user_input})
    st.session_state.state = app.invoke(st.session_state.state)
    st.rerun()

# After completion, show a meta summary at the bottom
if st.session_state.state.get("stage") == "completed":
    job = st.session_state.state.get("selected_job") or {}
    st.success("Application submitted and stored.")
    st.caption("Below is a summary of your submission.")
    st.json({
        "job": {"id": job.get("id"), "title": job.get("title")},
        "profile": st.session_state.state.get("profile", {}),
        # "answers_count": len(st.session_state.state.get("answers", {})),
        # "storage": APPLICATIONS_PATH,
    })
